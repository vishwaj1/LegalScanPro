from fastapi import FastAPI, UploadFile, File, HTTPException, Body
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
import os
import uuid
import tempfile
import re
import boto3
import json
from pydantic import BaseModel
from typing import Dict, Any
from dotenv import load_dotenv
import os

# Load environment variables from the correct path
load_dotenv('/Users/vishwanthreddyjakka/LegalScanPro/.env')

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://legal-scan-pro-pyk5.vercel.app",
        "http://localhost:3000",
        "https://www.legal-scan-pro-pyk5.vercel.app"  # (optional, for www)
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# AI setup removed - using predefined form instead

# Predefined form structure for SAFE agreements
PREDEFINED_FORM_FIELDS = [
    {"id": "company_legal_name", "label": "What is the full legal name of the company?", "placeholder": "{{COMPANY_NAME}}", "type": "text", "required": True},
    {"id": "company_state", "label": "In which state is the company incorporated?", "placeholder": "{{COMPANY_STATE}}", "type": "text", "required": True},
    {"id": "governing_law_state", "label": "Which state's laws will govern this agreement?", "placeholder": "{{GOVERNING_LAW}}", "type": "text", "required": True},
    {"id": "investor_name", "label": "What is the full legal name of the investor?", "placeholder": "{{INVESTOR_NAME}}", "type": "text", "required": True},
    {"id": "investor_title", "label": "What is the investor's professional title (if any)?", "placeholder": "{{INVESTOR_TITLE}}", "type": "text", "required": False},
    {"id": "investor_address", "label": "What is the investor's mailing address?", "placeholder": "{{INVESTOR_ADDRESS}}", "type": "textarea", "required": True},
    {"id": "investor_email", "label": "What is the investor's email address?", "placeholder": "{{INVESTOR_EMAIL}}", "type": "email", "required": True},
    {"id": "purchase_amount", "label": "What is the purchase amount in USD (e.g., 50,000)?", "placeholder": "{{PURCHASE_AMOUNT}}", "type": "number", "required": True},
    {"id": "execution_date", "label": "On what date will this SAFE be executed? (MM/DD/YYYY)", "placeholder": "{{EXECUTION_DATE}}", "type": "date", "required": True},
    {"id": "valuation_cap", "label": "What is the post-money valuation cap in USD (e.g., 5,000,000)?", "placeholder": "{{VALUATION_CAP}}", "type": "number", "required": True},
    {"id": "company_signatory_name", "label": "What is the name of the authorized company signatory?", "placeholder": "{{COMPANY_SIGNATORY_NAME}}", "type": "text", "required": True},
    {"id": "company_signatory_title", "label": "What is the title of the authorized company signatory?", "placeholder": "{{COMPANY_SIGNATORY_TITLE}}", "type": "text", "required": True},
    {"id": "company_signatory_address", "label": "What is the mailing address of the company signatory?", "placeholder": "{{COMPANY_SIGNATORY_ADDRESS}}", "type": "textarea", "required": True},
    {"id": "company_signatory_email", "label": "What is the email address of the company signatory?", "placeholder": "{{COMPANY_SIGNATORY_EMAIL}}", "type": "email", "required": True}
]

R2_ENDPOINT = os.getenv("R2_ENDPOINT")
#MINIO_ENDPOINT = "http://minio:9000"
BUCKET = os.getenv("R2_BUCKET_NAME")
# BUCKET = "legal-docs"

# Clean up R2_ENDPOINT to remove any path suffix
if R2_ENDPOINT and "/" in R2_ENDPOINT.split("//")[1]:
    R2_ENDPOINT = R2_ENDPOINT.split("/")[0] + "//" + R2_ENDPOINT.split("//")[1].split("/")[0]

s3 = boto3.client(
    "s3",
    endpoint_url=R2_ENDPOINT,
    aws_access_key_id=os.getenv("R2_ACCESS_KEY"),
    aws_secret_access_key=os.getenv("R2_SECRET_KEY"),
)

# s3 = boto3.client(
#     "s3",
#     endpoint_url=MINIO_ENDPOINT,
#     aws_access_key_id="minio",
#     aws_secret_access_key="minio123",
# )

@app.on_event("startup")
def create_bucket_if_needed():
    if not BUCKET:
        return
    try:
        s3.head_bucket(Bucket=BUCKET)
    except:
        s3.create_bucket(Bucket=BUCKET)

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    try:
        content = await file.read()
        s3.put_object(Bucket=BUCKET, Key=file.filename, Body=content)
        return {"message": "File uploaded successfully", "filename": file.filename}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# New endpoint to get the predefined form structure
@app.get("/form-fields")
async def get_form_fields():
    """Return the predefined form fields for SAFE agreement"""
    return {"fields": PREDEFINED_FORM_FIELDS}

# New endpoint to list available templates
@app.get("/templates")
async def list_templates():
    """List available document templates from R2"""
    try:
        if not BUCKET:
            raise HTTPException(status_code=500, detail="R2_BUCKET_NAME not configured")
        
        # List templates from the legal-scan-pro folder
        response = s3.list_objects_v2(Bucket=BUCKET, Prefix="legal-scan-pro/")
        templates = []
        if 'Contents' in response:
            for obj in response['Contents']:
                # Only include actual templates, exclude filled_docs folder
                if (obj['Key'].endswith('.docx') and 
                    not obj['Key'].startswith('legal-scan-pro/filled_docs/') and
                    not obj['Key'].endswith('_filled.docx')):
                    
                    # Extract just the filename from the full path
                    template_name = obj['Key'].split('/')[-1]
                    templates.append({
                        "name": obj['Key'],  # Full path for download
                        "display_name": template_name.replace('.docx', '').replace('_', ' ').title()
                    })
        return {"templates": templates}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to list templates: {e}")

# Data models for the new form-based approach
class FormData(BaseModel):
    company_legal_name: str
    company_state: str
    governing_law_state: str
    investor_name: str
    investor_title: str = ""
    investor_address: str
    investor_email: str
    purchase_amount: float
    execution_date: str
    valuation_cap: float
    company_signatory_name: str
    company_signatory_title: str
    company_signatory_address: str
    company_signatory_email: str

class DocumentFillRequest(BaseModel):
    template_name: str
    form_data: FormData


@app.post("/document/fill")
async def fill_document(request: DocumentFillRequest):
    """Fill a document template with user-provided form data"""
    try:
        # Download template from R2 using the full path
        template_key = request.template_name  # This now contains the full path
        try:
            template_obj = s3.get_object(Bucket=BUCKET, Key=template_key)
            template_content = template_obj['Body'].read()
        except Exception as e:
            raise HTTPException(status_code=404, detail=f"Template not found: {request.template_name}")
        
        # Save template temporarily
        temp_template_path = f"/tmp/{uuid.uuid4()}_template.docx"
        with open(temp_template_path, "wb") as f:
            f.write(template_content)
        
        # Load document
        doc = Document(temp_template_path)
        
        # Create replacement mapping from form data
        replacements = {
            "{{COMPANY_NAME}}": request.form_data.company_legal_name,
            "{{Company_Name}}": request.form_data.company_legal_name,  # Handle mixed case
            "{{COMPANY_STATE}}": request.form_data.company_state,
            "{{Company_State}}": request.form_data.company_state,  # Handle mixed case
            "{{GOVERNING_LAW}}": request.form_data.governing_law_state,
            "{{Governing_Law}}": request.form_data.governing_law_state,  # Handle mixed case
            "{{INVESTOR_NAME}}": request.form_data.investor_name,
            "{{Investor_Name}}": request.form_data.investor_name,  # Handle mixed case
            "{{INVESTOR_TITLE}}": request.form_data.investor_title,
            "{{Investor_Title}}": request.form_data.investor_title,  # Handle mixed case
            "{{INVESTOR_ADDRESS}}": request.form_data.investor_address,
            "{{Investor_Address}}": request.form_data.investor_address,  # Handle mixed case
            "{{INVESTOR_EMAIL}}": request.form_data.investor_email,
            "{{Investor_Email}}": request.form_data.investor_email,  # Handle mixed case
            "{{PURCHASE_AMOUNT}}": f"${request.form_data.purchase_amount:,.2f}",
            "{{Purchase_Amount}}": f"${request.form_data.purchase_amount:,.2f}",  # Handle mixed case
            "{{EXECUTION_DATE}}": request.form_data.execution_date,
            "{{Execution_Date}}": request.form_data.execution_date,  # Handle mixed case
            "{{VALUATION_CAP}}": f"${request.form_data.valuation_cap:,.2f}",
            "{{Valuation_Cap}}": f"${request.form_data.valuation_cap:,.2f}",  # Handle mixed case
            "{{COMPANY_SIGNATORY_NAME}}": request.form_data.company_signatory_name,
            "{{Company_Signatory_Name}}": request.form_data.company_signatory_name,  # Handle mixed case
            "{{COMPANY_SIGNATORY_TITLE}}": request.form_data.company_signatory_title,
            "{{Company_Signatory_Title}}": request.form_data.company_signatory_title,  # Handle mixed case
            "{{COMPANY_SIGNATORY_ADDRESS}}": request.form_data.company_signatory_address,
            "{{Company_Signatory_Address}}": request.form_data.company_signatory_address,  # Handle mixed case
            "{{COMPANY_SIGNATORY_EMAIL}}": request.form_data.company_signatory_email,
            "{{Company_Signatory_Email}}": request.form_data.company_signatory_email  # Handle mixed case
        }
        

        
        # Replace placeholders in all paragraphs (handling formatted text)
        for paragraph in doc.paragraphs:
            for placeholder, replacement in replacements.items():
                if placeholder in paragraph.text:
                    # First try: Replace in individual runs
                    replaced_in_runs = False
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement)
                            replaced_in_runs = True
                    
                    # Second try: If not replaced in runs, replace at paragraph level
                    if not replaced_in_runs and placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, replacement)
        
        # Replace placeholders in headers and footers
        for section in doc.sections:
            # Primary header
            if section.header:
                for header in section.header.paragraphs:
                    for placeholder, replacement in replacements.items():
                        if placeholder in header.text:
                            # First try: Replace in individual runs
                            for run in header.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, replacement)
                            # Also try at paragraph level
                            if placeholder in header.text:
                                header.text = header.text.replace(placeholder, replacement)
            
            # First page header
            if section.different_first_page_header_footer and section.first_page_header:
                for header in section.first_page_header.paragraphs:
                    for placeholder, replacement in replacements.items():
                        if placeholder in header.text:
                            print(f"Found placeholder {placeholder} in first page header: {header.text[:100]}...")
                            header.text = header.text.replace(placeholder, replacement)
            
            # Primary footer
            if section.footer:
                for footer in section.footer.paragraphs:
                    for placeholder, replacement in replacements.items():
                        if placeholder in footer.text:
                            print(f"Found placeholder {placeholder} in primary footer: {footer.text[:100]}...")
                            footer.text = footer.text.replace(placeholder, replacement)
            
            # First page footer
            if section.different_first_page_header_footer and section.first_page_footer:
                for footer in section.first_page_footer.paragraphs:
                    for placeholder, replacement in replacements.items():
                        if placeholder in footer.text:
                            print(f"Found placeholder {placeholder} in first page footer: {footer.text[:100]}...")
                            footer.text = footer.text.replace(placeholder, replacement)
        
        # Replace placeholders in tables (if any)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, replacement in replacements.items():
                            if placeholder in paragraph.text:
                                # First try: Replace in individual runs
                                replaced_in_runs = False
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, replacement)
                                        replaced_in_runs = True
                                
                                # Second try: If not replaced in runs, replace at paragraph level
                                if not replaced_in_runs and placeholder in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder, replacement)
        
        # Save filled document
        filled_filename = f"{uuid.uuid4()}_filled.docx"
        filled_path = f"/tmp/{filled_filename}"
        doc.save(filled_path)
        
        # Upload to R2
        key = f"filled_docs/{filled_filename}"
        with open(filled_path, "rb") as f:
            s3.put_object(
                Bucket=BUCKET,
                Key=key,
                Body=f,
                ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        # Generate public URL
        base_url = os.getenv("R2_PUBLIC_BASE_URL")
        if not base_url:
            raise HTTPException(status_code=500, detail="R2_PUBLIC_BASE_URL not configured")
        public_url = f"{base_url}/{key}"
        
        # Clean up temp files
        os.remove(temp_template_path)
        os.remove(filled_path)
        
        return {
            "success": True,
            "public_preview_url": public_url,
            "download_url": public_url
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fill document: {str(e)}")
